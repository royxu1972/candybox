from docx import Document
from docx.shared import Inches
import os

PATH = "D:\\Work CT Application\\report\\"

'''
convert the table in a.docx file to a.html
'''
def convert(filename):
	document = Document(PATH + filename)
	output = "<table class='table table-hover table-layout: automatic' width='100%'>\n"
	
	content = document.tables[0].cell(0,1).text
	content = content.replace('\n','<br>')
	output += "<tr><th width='17%'>Name</th><td width='83%'>" + content + "</td></tr>\n"

	for row in range(1,16):
		item = document.tables[0].cell(row,0).text
		content = document.tables[0].cell(row,1).text
		content = content.replace('\n','<br>')
		output += "<tr><th>" + item + "</th><td>" + content + "</td></tr>\n"
	output += "</table>"

	# write
	tp = filename[0:filename.find('.')].split()
	name = tp[0] + tp[1]
	f1 = open( "D:\\Workspace\\DocApplication\\docs\\" + name + '.html','w')
	f1.write(output)
	f1.close
	return name

'''
main
'''
list_dirs = os.walk(PATH)
outlist = ""
for root, dirs, files in list_dirs: 
	for f in files:
		print(f) 
		ids = convert(f)
		name = f[f.find(' ')+1:f.find('.')]
		outlist += "<li id='" + ids + "' class='app_list'><a href='#'>" + name + "</a></li>\n"
f = open( "list.html",'w')
f.write(outlist)
f.close