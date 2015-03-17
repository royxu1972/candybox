from docx import Document
from docx.shared import Inches
import os

filename = "D:\\Work CT Application\\software list.docx"
document = Document(filename)



def DoTable(table):
	output = "<table class='table table-hover table-layout: automatic' width='100%'>\n"
	output += "<tr><td width='28%'></td><td width='36'></td><th width='10%'>Input</th><th width='13%'>Configuration</th><th width='13%'>Event</th></tr>\n"

	columns = table.column_cells(0)
	index = 1

	while( index < len(columns) ):
		item = columns[index].text

		# main item, colspan
		if( index + 1 < len(columns) and item != columns[index+1].text ):
			print(item)
			output += "<tr><th colspan=5><h4><em>" + item + "</em></h4></th></tr>\n"
			index += 1

		else:
			num = item[item.find('(')+1: item.find(')')]
			print(num)
			output += "<tr><td rowspan=" + num + ">" + item + "</td>"
			for c in range(1,5):
				output += "<td>" + table.cell(index,c).text + "</td>"
			output += "</tr>\n"
			index += 1

			for r in range(0,int(num)-1):
				output += "<tr>"
				for c in range(1,5):
					output += "<td>" + table.cell(index,c).text + "</td>"
				index += 1
				output += "</tr>\n"

	
	output += "</table>"
	output = output.replace("√","&radic;")

	return output


out = "<div class='panel panel-default'>\n"
out += "<div class='panel-heading'><h3>Application Software(92)</h3></div>\n"
out += DoTable(document.tables[0])
out += "</div><br>\n\n"

out += "<div class='panel panel-default'>\n"
out += "<div class='panel-heading'><h3>System Software(36)</h3></div>\n"
out += DoTable(document.tables[1])
out += "</div><br>\n\n"

out += "<div class='panel panel-default'>\n"
out += "<div class='panel-heading'><h3>Web Application (32)</h3></div>\n"
out += DoTable(document.tables[2])
out += "</div><br>\n\n"

out += "<div class='panel panel-default'>\n"
out += "<div class='panel-heading'><h3>Mobile Application (40)</h3></div>\n"
out += DoTable(document.tables[3])

f1 = open('D:\\Workspace\\WebReportDoc\\real.html','w')
f1.write(out)
f1.close