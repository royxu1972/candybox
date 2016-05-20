from docx import Document
from docx.shared import Inches
import os
import re
import sys

PATH = "D:\\Work CT Application\\report\\"

'''
convert the table in a.docx file to a.html
'''
def convert(filename):
	document = Document(PATH + filename)
	output = "<table class='table table-hover table-layout: automatic' width='100%'>\n"
	
	content = document.tables[0].cell(0,1).text
	content = content.replace('\n','<br>')
	output += "<tr><th width='17%'>Name</th><td width='83%'>" + content + "</td></tr>\n"

	for row in range(1,16):
		item = document.tables[0].cell(row,0).text
		content = document.tables[0].cell(row,1).text
		content = content.replace('\n','<br>')
		output += "<tr><th>" + item + "</th><td>" + content + "</td></tr>\n"
	output += "</table>"

	# replace "^"
	strinfo = re.compile('\^\d+\s*')
	c = strinfo.findall(output)
	for i in range(0, len(c)):
		c[i] = c[i][1:]
		c[i] = c[i].strip()
	for i in range(0, len(c)):
		output = strinfo.sub('<sup>' + c[i] + '</sup>', output, 1)

	# write
	tp = filename[0:filename.find('.')].split()
	name = tp[0] + tp[1]
	f1 = open( "D:\\Workspace\\WebReportDoc\\docs\\" + name + '.html','w')
	f1.write(output)
	f1.close
	return name


def appendcontent():
	# SPL
	f = open('1.txt')
	all_text = f.read()
	f.close()
	f = open( "D:\\Workspace\\WebReportDoc\\docs\\2006SPL.html",'a')
	f.write(all_text)
	f.close

	# real cases
	f = open('2.txt')
	all_text = f.read()
	f.close()
	f = open( "D:\\Workspace\\WebReportDoc\\docs\\2011real.html",'a')
	f.write(all_text)
	f.close





'''
t = "posdad 5^1 2^128 3^4, ddd 9^10 2^1"
strinfo = re.compile('\^\d+\s*')
c = strinfo.findall(t)
for i in range(0, len(c)):
	c[i] = c[i][1:]
	c[i] = c[i].strip()
print(c)

for i in range(0, len(c)):
	t = strinfo.sub('<sup>' + c[i] + '</sup>', t, 1)
print(t)
sys.exit(0)
'''

'''
main
'''
list_dirs = os.walk(PATH)
outlist = ""
for root, dirs, files in list_dirs: 
	for f in files:
		print(f) 
		ids = convert(f)
		name = f[f.find(' ')+1:f.find('.')]
		outlist += "<li id='" + ids + "' class='app_list'><a href='#'>" + name + "</a></li>\n"
f = open( "list.html",'w')
f.write(outlist)
f.close
appendcontent()
